"""
Export utilities for generating PDF and DOCX files from job descriptions
"""
from pathlib import Path
import html
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_pdf(content: str, output_path: Path) -> None:
    """Generate a PDF from job description text"""
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )

    # Define styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=12,
        alignment=0  # Left align (TA_LEFT = 0)
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=8,
        spaceBefore=12
    )
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        spaceAfter=8
    )

    story = []
    lines = content.split('\n')
    skip_until_next_section = False
    start_rendering = False

    for line in lines:
        line = line.strip()

        # Skip separator lines
        if line.startswith('===') or line.startswith('---'):
            continue

        # Start rendering from JOB INFORMATION section
        if line.upper() == 'JOB INFORMATION':
            start_rendering = True

        # Skip everything before JOB INFORMATION
        if not start_rendering:
            continue

        # Skip Classification Job Information section entirely
        if line.upper() == 'CLASSIFICATION JOB INFORMATION':
            skip_until_next_section = True
            continue

        # Check if we've reached the next major section (ALL CAPS line)
        if skip_until_next_section and line == line.upper() and len(line) > 0 and not line.startswith('•'):
            skip_until_next_section = False

        # Skip lines within Classification Job Information section
        if skip_until_next_section:
            continue

        # Skip Exclusion Status line
        if line.startswith('Exclusion Status:'):
            continue

        # Escape HTML special characters for ReportLab
        escaped_line = html.escape(line)

        # Detect major section headings (ALL CAPS, no colon, standalone)
        if line == line.upper() and len(line) > 0 and not line.startswith('•') and ':' not in line:
            # Convert to Title Case
            title_case = line.title()
            story.append(Paragraph(html.escape(title_case), title_style))
        # Detect subsection headings (ends with colon, short line)
        elif line.endswith(':') and len(line) < 50:
            story.append(Paragraph(escaped_line, heading_style))
        # Regular text or bullets
        elif line:
            story.append(Paragraph(escaped_line, normal_style))
        # Empty line
        else:
            story.append(Spacer(1, 0.1 * inch))

    doc.build(story)


def generate_docx(content: str, output_path: Path) -> None:
    """Generate a DOCX from job description text"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = content.split('\n')
    skip_until_next_section = False
    start_rendering = False

    for line in lines:
        line = line.strip()

        # Skip separator lines
        if line.startswith('===') or line.startswith('---'):
            continue

        # Start rendering from JOB INFORMATION section
        if line.upper() == 'JOB INFORMATION':
            start_rendering = True

        # Skip everything before JOB INFORMATION
        if not start_rendering:
            continue

        # Skip Classification Job Information section entirely
        if line.upper() == 'CLASSIFICATION JOB INFORMATION':
            skip_until_next_section = True
            continue

        # Check if we've reached the next major section (ALL CAPS line)
        if skip_until_next_section and line == line.upper() and len(line) > 0 and not line.startswith('•'):
            skip_until_next_section = False

        # Skip lines within Classification Job Information section
        if skip_until_next_section:
            continue

        # Skip Exclusion Status line
        if line.startswith('Exclusion Status:'):
            continue

        # Detect major section headings (ALL CAPS, no colon, standalone)
        if line == line.upper() and len(line) > 0 and not line.startswith('•') and ':' not in line:
            # Convert to Title Case
            title_case = line.title()
            p = doc.add_heading(title_case, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Detect subsection headings (ends with colon, short line)
        elif line.endswith(':') and len(line) < 50:
            doc.add_heading(line, level=2)
        # Bullet points
        elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
            text = line[1:].strip()
            doc.add_paragraph(text, style='List Bullet')
        # Regular text (including field: value pairs)
        elif line:
            doc.add_paragraph(line)
        # Empty line
        else:
            doc.add_paragraph('')

    doc.save(str(output_path))
