"""DOCX processing utilities for position descriptions."""
from pathlib import Path

from docx import Document


class DOCXProcessor:
    """Extract text from DOCX position descriptions."""

    def __init__(self, docx_path: str | Path):
        """Initialize processor with DOCX path."""
        self.docx_path = Path(docx_path)
        if not self.docx_path.exists():
            raise FileNotFoundError(f"DOCX not found: {self.docx_path}")

    def extract_text(self) -> str:
        """
        Extract all text from DOCX.

        Returns:
            Full text content from all paragraphs
        """
        doc = Document(self.docx_path)

        # Extract text from paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        # Combine all text
        all_text = "\n\n".join(paragraphs)
        if table_texts:
            all_text += "\n\n--- Tables ---\n" + "\n".join(table_texts)

        return all_text

    def extract_metadata(self) -> dict[str, str | None]:
        """
        Extract DOCX metadata.

        Returns:
            Dictionary with title, author, subject, etc.
        """
        doc = Document(self.docx_path)
        core_props = doc.core_properties

        return {
            "title": core_props.title,
            "author": core_props.author,
            "subject": core_props.subject,
            "keywords": core_props.keywords,
            "comments": core_props.comments,
            "created": str(core_props.created) if core_props.created else None,
            "modified": str(core_props.modified) if core_props.modified else None,
        }
