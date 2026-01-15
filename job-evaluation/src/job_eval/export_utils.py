"""
Export utilities for generating PDF and DOCX files from job evaluation results
"""
import html
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def generate_pdf(content: str, output_path: Path) -> None:
    """Generate a PDF from job evaluation text"""
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72,
    )

    # Define styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=16,
        spaceAfter=12,
        alignment=TA_CENTER,
    )
    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading2"],
        fontSize=14,
        spaceAfter=8,
        spaceBefore=12,
    )
    subheading_style = ParagraphStyle(
        "CustomSubheading",
        parent=styles["Heading3"],
        fontSize=12,
        spaceAfter=6,
        spaceBefore=8,
    )
    normal_style = ParagraphStyle(
        "CustomNormal",
        parent=styles["Normal"],
        fontSize=11,
        leading=16,
        spaceAfter=8,
    )

    story = []
    lines = content.split("\n")

    for line in lines:
        # Escape HTML special characters for ReportLab
        escaped_line = html.escape(line)

        # Detect main title (first line)
        if "JOB EVALUATION ANALYSIS REPORT" in line:
            story.append(Paragraph(escaped_line, title_style))
        # Detect major section headers (TOOL 1:, TOOL 2:, etc.)
        elif line.startswith("TOOL ") and ":" in line:
            story.append(Paragraph(escaped_line, heading_style))
        # Detect separator lines
        elif line.strip().startswith("===") or line.strip().startswith("---"):
            story.append(Spacer(1, 0.1 * inch))
        # Detect subsection headers (ends with colon, not indented, short)
        elif line.endswith(":") and not line.startswith(" ") and len(line) < 50:
            story.append(Paragraph(escaped_line, subheading_style))
        # Regular text or bullets
        elif line.strip():
            story.append(Paragraph(escaped_line, normal_style))
        # Empty line
        else:
            story.append(Spacer(1, 0.1 * inch))

    doc.build(story)


def generate_docx(content: str, output_path: Path) -> None:
    """Generate a DOCX from job evaluation text"""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = content.split("\n")

    for line in lines:
        # Detect main title (first line)
        if "JOB EVALUATION ANALYSIS REPORT" in line:
            p = doc.add_heading(line, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Detect major section headers (TOOL 1:, TOOL 2:, etc.)
        elif line.startswith("TOOL ") and ":" in line:
            doc.add_heading(line, level=1)
        # Skip separator lines
        elif line.strip().startswith("===") or line.strip().startswith("---"):
            continue
        # Detect subsection headers (ends with colon, not indented, short)
        elif line.endswith(":") and not line.startswith(" ") and len(line) < 50:
            doc.add_heading(line, level=2)
        # Bullet points (detect by leading spaces and bullet)
        elif line.strip().startswith("•"):
            # Count leading spaces to determine indentation level
            stripped = line.lstrip()
            indent_level = (len(line) - len(stripped)) // 2
            text = stripped[1:].strip()  # Remove bullet character
            p = doc.add_paragraph(text, style="List Bullet")
            # Adjust indentation if needed
            if indent_level > 0:
                p.paragraph_format.left_indent = Pt(indent_level * 20)
        # Regular text
        elif line.strip():
            doc.add_paragraph(line)
        # Empty line
        else:
            doc.add_paragraph("")

    doc.save(str(output_path))
